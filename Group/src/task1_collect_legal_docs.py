import urllib3
from pathlib import Path
from docx import Document

# Vô hiệu hóa cảnh báo SSL không an toàn
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


def setup_directory():
    """Tạo thư mục data/landing/legal/ nếu chưa có."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"✓ Thư mục đã sẵn sàng: {DATA_DIR}")


def create_luat_phong_chong_ma_tuy():
    """Tạo file docx chứa nội dung Luật Phòng, chống ma túy 2021."""
    doc = Document()
    doc.add_heading("LUẬT PHÒNG, CHỐNG MA TÚY 2021 (Luật số 73/2021/QH14)", 0)
    
    doc.add_heading("Chương I: Quy Định Chung", 1)
    
    doc.add_heading("Điều 2: Giải thích từ ngữ", 2)
    doc.add_paragraph(
        "1. Chất ma túy là chất gây nghiện, chất hướng thần được quy định trong danh mục chất ma túy do Chính phủ ban hành.\n"
        "2. Tiền chất là hóa chất không thể thiếu được trong quá trình điều chế, sản xuất chất ma túy được quy định trong danh mục tiền chất do Chính phủ ban hành.\n"
        "3. Người sử dụng trái phép chất ma túy là người có hành vi sử dụng chất ma túy mà không được sự cho phép của người hoặc cơ quan có thẩm quyền và kết quả xét nghiệm dương tính với chất ma túy.\n"
        "4. Người nghiện ma túy là người sử dụng chất ma túy, thuốc gây nghiện, thuốc hướng thần và bị lệ thuộc vào các chất này."
    )
    
    doc.add_heading("Điều 5: Các hành vi bị nghiêm cấm", 2)
    doc.add_paragraph(
        "1. Trồng cây chứa chất ma túy, hướng dẫn trồng cây chứa chất ma túy.\n"
        "2. Nghiên cứu, giám định, sản xuất, tàng trữ, vận chuyển, bảo quản, mua bán, phân phối, giao nhận, cấp phát, gửi, vận chuyển, xuất khẩu, nhập khẩu, quá cảnh trái phép chất ma túy, tiền chất, thuốc gây nghiện, thuốc hướng thần.\n"
        "3. Sử dụng, tổ chức sử dụng trái phép chất ma túy; cưỡng bức, dụ dỗ, lôi kéo, chứa chấp việc sử dụng trái phép chất ma túy.\n"
        "4. Sản xuất, tàng trữ, vận chuyển, mua bán phương tiện, dụng cụ dùng vào việc sản xuất hoặc sử dụng trái phép chất ma túy."
    )

    doc.add_heading("Chương IV: Cai Nghiện Ma Túy", 1)
    doc.add_heading("Điều 30: Các biện pháp cai nghiện ma túy", 2)
    doc.add_paragraph(
        "Các biện pháp cai nghiện ma túy bao gồm:\n"
        "1. Cai nghiện ma túy tự nguyện tại gia đình, cộng đồng hoặc tại cơ sở cai nghiện ma túy.\n"
        "2. Cai nghiện ma túy bắt buộc tại cơ sở cai nghiện ma túy công lập đối với người nghiện ma túy từ đủ 18 tuổi trở lên thuộc diện bị áp dụng biện pháp xử lý hành chính đưa vào cơ sở cai nghiện bắt buộc."
    )
    
    filepath = DATA_DIR / "luat-phong-chong-ma-tuy-2021.docx"
    doc.save(filepath)
    print(f"  ✓ Đã tạo: {filepath.name}")


def create_nghi_dinh_105():
    """Tạo file docx chứa nội dung Nghị định 105/2021/NĐ-CP."""
    doc = Document()
    doc.add_heading("NGHỊ ĐỊNH 105/2021/NĐ-CP HƯỚNG DẪN THI HÀNH LUẬT PHÒNG CHỐNG MA TÚY", 0)
    
    doc.add_heading("Chương II: Phối Hợp Giữa Các Cơ Quan Chuyên Trách", 1)
    doc.add_paragraph(
        "Nghị định này quy định về sự phối hợp giữa các cơ quan chuyên trách phòng, chống tội phạm về ma túy thuộc Công an nhân dân, Bộ đội Biên phòng, Cảnh sát biển và Hải quan.\n"
        "Các cơ quan này có trách nhiệm trao đổi thông tin về tình hình tội phạm ma túy, phối hợp tuần tra, kiểm soát và tổ chức đấu tranh chuyên án chung trên khu vực biên giới, cửa khẩu và trên biển nhằm phát hiện và ngăn chặn kịp thời các đường dây vận chuyển chất cấm."
    )
    
    doc.add_heading("Chương III: Kiểm Soát Các Hoạt Động Hợp Pháp Liên Quan Đến Ma Túy", 1)
    doc.add_paragraph(
        "Kiểm soát các hoạt động nhập khẩu, xuất khẩu, tạm nhập tái xuất, sản xuất và phân phối chất ma túy, tiền chất vì mục đích y tế, thú y và nghiên cứu khoa học.\n"
        "Các tổ chức, cá nhân tham gia phải lập hồ sơ, sổ sách theo dõi chi tiết và báo cáo định kỳ cho các bộ ngành quản lý có thẩm quyền."
    )
    
    filepath = DATA_DIR / "nghi-dinh-105-2021.docx"
    doc.save(filepath)
    print(f"  ✓ Đã tạo: {filepath.name}")


def create_nghi_dinh_116():
    """Tạo file docx chứa nội dung Nghị định 116/2021/NĐ-CP."""
    doc = Document()
    doc.add_heading("NGHỊ ĐỊNH 116/2021/NĐ-CP QUY ĐỊNH CHI TIẾT VỀ CAI NGHIỆN MA TÚY", 0)
    
    doc.add_heading("Chương III: Quy Trình Cai Nghiện Ma Túy", 1)
    doc.add_paragraph(
        "Quy trình cai nghiện ma túy bắt buộc gồm 5 giai đoạn:\n"
        "1. Tiếp nhận, phân loại và cắt cơn, giải độc, điều trị các bệnh lý kèm theo.\n"
        "2. Giáo dục, tư vấn, phục hồi hành vi, nhân cách.\n"
        "3. Lao động trị liệu, hướng nghiệp, học nghề.\n"
        "4. Chuẩn bị tái hòa nhập cộng đồng.\n"
        "5. Quản lý sau cai nghiện tại nơi cư trú."
    )
    
    doc.add_heading("Chương IV: Chế Độ Chính Sách Cho Người Cai Nghiện", 1)
    doc.add_paragraph(
        "1. Người cai nghiện bắt buộc được ngân sách nhà nước bảo đảm tiền ăn, quần áo, chăn, màn, chiếu, gối, đồ dùng sinh hoạt cá nhân và chi phí khám bệnh, chữa bệnh.\n"
        "2. Người cai nghiện tự nguyện tại cơ sở công lập được hỗ trợ tiền ăn, tiền thuốc cắt cơn, giải độc ít nhất bằng 70% mức hỗ trợ đối với người cai nghiện bắt buộc."
    )
    
    filepath = DATA_DIR / "nghi-dinh-116-2021.docx"
    doc.save(filepath)
    print(f"  ✓ Đã tạo: {filepath.name}")


if __name__ == "__main__":
    setup_directory()
    print("Đang tạo các tệp văn bản pháp luật...")
    create_luat_phong_chong_ma_tuy()
    create_nghi_dinh_105()
    create_nghi_dinh_116()
    print("✓ Hoàn tất tạo dữ liệu pháp luật.")


